import argparse
import os
import pandas as pd
from docx import Document
from docx.shared import Inches
import datetime
from PIL import Image
import io
import glob

class TestCaseGenerator:
    def __init__(self, excel_path, template_path, input_folder, output_folder):
        """
        Initialize the test case generator.
        
        Args:
            excel_path (str): Path to the Excel file containing test case information
            template_path (str): Path to the Word template file
            input_folder (str): Path to the input folder containing test case folders
            output_folder (str): Path to the output folder for generated documents
        """
        self.excel_path = excel_path
        self.template_path = template_path
        self.input_folder = input_folder
        self.output_folder = output_folder
        
        # Create output folder if it doesn't exist
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
            
        # Load test case data from Excel
        self.test_cases = self._load_test_cases()
        
    def _load_test_cases(self):
        """Load test case data from Excel file."""
        try:
            df = pd.read_excel(self.excel_path)
            # Assuming the Excel columns match these names
            required_columns = ['案例編號', '交易代號', '交易名稱', '情境說明']
            
            # Verify all required columns exist
            missing_columns = [col for col in required_columns if col not in df.columns]
            if missing_columns:
                raise ValueError(f"Missing required columns in Excel: {', '.join(missing_columns)}")
                
            # Convert DataFrame to dictionary for easier lookup
            test_cases = {}
            for _, row in df.iterrows():
                try:
                    completion_date = pd.to_datetime(row['實際完成日期']).strftime('%Y-%m-%d')
                except:
                    completion_date = row['實際完成日期']
                    
                test_cases[row['案例編號']] = {
                    '交易代號': row['交易代號'],
                    '交易名稱': row['交易名稱'],
                    '情境說明': row['情境說明'],
                    '實際完成日期': completion_date
                }
            return test_cases
            
        except Exception as e:
            print(f"Error loading Excel file: {e}")
            return {}
        
    ##讀 input 下的 測案資料夾， 需要出現在 EXCEL 中
    def _get_case_folders(self):
        """Get list of test case folders in the input directory."""
        return [f for f in os.listdir(self.input_folder) 
                if os.path.isdir(os.path.join(self.input_folder, f)) 
                and f in self.test_cases]
    
    def _get_images(self, case_folder):
        """Get list of image files in the case's image folder, organized by platform."""
        case_path = os.path.join(self.input_folder, case_folder)
        platforms = ['web', 'ios', 'android']
        
        image_files_by_platform = {}
        
        for platform in platforms:
            platform_path = os.path.join(case_path, platform)
            if os.path.exists(platform_path):
                image_files = []
                for ext in ['*.png', '*.jpg', '*.jpeg', '*.gif', '*.bmp']:
                    image_files.extend(glob.glob(os.path.join(platform_path, ext)))
                if image_files:
                    image_files_by_platform[platform] = sorted(image_files)
                    
        if not image_files_by_platform:
            # Fall back to checking old 'image' folder structure
            image_folder = os.path.join(case_path, 'image')
            if os.path.exists(image_folder):
                image_files = []
                for ext in ['*.png', '*.jpg', '*.jpeg', '*.gif', '*.bmp']:
                    image_files.extend(glob.glob(os.path.join(image_folder, ext)))
                if image_files:
                    image_files_by_platform['default'] = sorted(image_files)
            else:
                # Try reading images directly from case folder
                image_files = []
                for ext in ['*.png', '*.jpg', '*.jpeg', '*.gif', '*.bmp']:
                    image_files.extend(glob.glob(os.path.join(case_path, ext)))
                if image_files:
                    image_files_by_platform['default'] = sorted(image_files)
                else:
                    print(f"Warning: No images found for {case_folder}")
                
        return image_files_by_platform
    
    def generate_documents(self):
        """Generate test case documents for all case folders."""
        case_folders = self._get_case_folders()
        
        if not case_folders:
            print("No matching test case folders found in input directory.")
            return
            
        print(f"Found {len(case_folders)} test case folders to process.")
        
        for case_folder in case_folders:
            self.generate_document(case_folder)
            
        print(f"Completed generating {len(case_folders)} test case documents.")
    
    def _resize_image_if_needed(self, img_path, max_width_inches=5.0, max_height_inches=3.5):
        """
        Resize image if it's too large for the Word document.
        
        Args:
            img_path (str): Path to the image file
            max_width_inches (float): Maximum width in inches
            max_height_inches (float): Maximum height in inches
            
        Returns:
            tuple: (image_path, width_inches, height_inches)
        """
        try:
            # Open the image using PIL
            with Image.open(img_path) as img:
                # Get original dimensions
                width, height = img.size
                
                # Convert to inches (assuming 96 DPI)
                dpi = 96
                width_inches = width / dpi
                height_inches = height / dpi
                
                # Check if image needs resizing
                if width_inches > max_width_inches or height_inches > max_height_inches:
                    # Calculate scale factor
                    width_scale = max_width_inches / width_inches
                    height_scale = max_height_inches / height_inches
                    scale = min(width_scale, height_scale)
                    
                    # Calculate new dimensions
                    new_width_inches = width_inches * scale
                    new_height_inches = height_inches * scale
                    
                    return img_path, new_width_inches, new_height_inches
                    
            # If image is already small enough
            return img_path, width_inches, height_inches
            
        except Exception as e:
            print(f"Error processing image {img_path}: {e}")
            # Return default values
            return img_path, 4.0, 2.5
    
    def generate_document(self, case_id):
        """
        Generate a single test case document.
        
        Args:
            case_id (str): Test case ID/folder name
        """
        print(f"Processing test case: {case_id}")
        
        # Check if test case exists in Excel data
        if case_id not in self.test_cases:
            print(f"Warning: Test case {case_id} not found in Excel data. Skipping.")
            return
            
        # Get case data
        case_data = self.test_cases[case_id]
        
        # Load Word template
        doc = Document(self.template_path)
        
        # Find the table in the template
        if not doc.tables:
            print(f"Error: No table found in Word template for {case_id}")
            return
            
        table = doc.tables[0]  # Assume first table is the test case table
        
        # Fill in the table cells
        # Assume the table has cells for each required field
        
        # Map data to fields in the Word table
        # Assuming a simple 5x2 table with headers in first column and data in second column
        field_map = {
            "測試日期": f"{case_data['實際完成日期']}",
            "測試案例": f"{case_id}_{case_data['情境說明']}",
            "測試環境": "",  # Could be parameterized if needed
            #"測試資料": f"交易代號: {case_data['交易代號']}, 交易名稱: {case_data['交易名稱']}",
            "測試資料": "",
            "測試結果": ""  # Will add images here
        }
        
        # Fill in the table
        for row in table.rows:
            if len(row.cells) >= 2:  # Ensure there are at least 2 cells
                header = row.cells[0].text.strip()
                if header in field_map:
                    row.cells[1].text = field_map[header]
        
        # Add images to the test result cell
        image_files_by_platform = self._get_images(case_id)
        if image_files_by_platform:
            print(image_files_by_platform)
            # Find the row with test results
            for row in table.rows:
                if len(row.cells) >= 2 and row.cells[0].text.strip() == "測試結果":
                    result_cell = row.cells[1]
                    # Clear any existing text
                    result_cell.text = ""
                    
                    # Add each image
                    paragraph = result_cell.paragraphs[0]
                    count=0
                    for platform,images in image_files_by_platform.items():
                        run = paragraph.add_run()
                        if (len(images)!=0):
                            if (count!=0):
                                run.add_break()
                            run.add_text(f"{platform}:")
                            run.add_break()

                        for img_path in images:
                            img_path, width_inches, height_inches = self._resize_image_if_needed(img_path)
                            
                            run.add_picture(img_path, width=Inches(width_inches), height=Inches(height_inches))
                            
                            # Add image caption
                            caption_run = run.add_break()
                            #run.add_text(f"圖片: {os.path.basename(img_path)}")
                            
                            # Add a line break between images
                            if img_path != images[-1]:  # Don't add break after last image
                                run.add_break()
                                run.add_break()  # Add extra space between images
                        count+=1
                        
        else:
            print(f"Warning: No images found for test case {case_id}")
        
        # Save the document
        output_path = os.path.join(self.output_folder, f"{case_id}_test_report.docx")
        doc.save(output_path)
        print(f"Generated test report: {output_path}")

def main():
    """Main function to run the test case generator."""
    # Configuration
    excel_path = "./createFile/template/test_case_template.xlsx"  # Path to Excel file
    template_path = "./createFile/template/test_report_template.docx"  # Path to Word template
    input_folder = "./createFile/input"  # Input folder containing test case folders
    output_folder = "./createFile/output/"  # Output folder for generated documents


    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=str, default=input_folder, help="images path default:./createFile/input/")
    parser.add_argument("--output", type=str, default=output_folder, help="output test doc path default:./createFile/output/")
    parser.add_argument("--excelpath", type=str, default=excel_path, help="excel path")
    parser.add_argument("--wordtmp", type=str, default=template_path, help="word template path")
    args = parser.parse_args()
    # Create generator instance
    generator = TestCaseGenerator(args.excelpath, args.wordtmp, args.input, args.output)
    
    # Generate documents
    generator.generate_documents()
    
    print("Test case document generation complete!")

if __name__ == "__main__":
    main()